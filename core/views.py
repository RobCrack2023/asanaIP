from django.utils import timezone
from rest_framework import viewsets, status
from rest_framework.decorators import action, api_view, permission_classes
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.response import Response
from django.contrib.auth import authenticate, login, logout

from rest_framework.permissions import BasePermission
from .models import User, Area, Team, Project, Section, Task, Asset, Organization, Plan, Notification
from .serializers import (
    UserSerializer, AreaSerializer, TeamSerializer,
    ProjectSerializer, ProjectListSerializer, SectionSerializer, TaskSerializer,
    AssetSerializer, OrganizationSerializer, PlanSerializer, NotificationSerializer,
)
from . import emails
from . import notifications as notif


class IsSuperAdmin(BasePermission):
    def has_permission(self, request, view):
        return request.user.is_authenticated and request.user.is_super_admin


@api_view(['POST'])
@permission_classes([AllowAny])
def login_view(request):
    username = request.data.get('username')
    password = request.data.get('password')
    user = authenticate(request, username=username, password=password)
    if user is None:
        return Response({'error': 'Credenciales inválidas'}, status=status.HTTP_401_UNAUTHORIZED)
    login(request, user)
    return Response(UserSerializer(user).data)


@api_view(['POST'])
def logout_view(request):
    logout(request)
    return Response({'ok': True})


@api_view(['GET'])
def me_view(request):
    data = UserSerializer(request.user).data
    if request.user.organization:
        data['organization_name'] = request.user.organization.name
        data['organization_slug'] = request.user.organization.slug
    return Response(data)


class UserViewSet(viewsets.ModelViewSet):
    queryset = User.objects.all()
    serializer_class = UserSerializer

    def get_queryset(self):
        qs = super().get_queryset()
        user = self.request.user
        if not user.is_super_admin:
            qs = qs.filter(organization=user.organization)
        return qs


class AreaViewSet(viewsets.ModelViewSet):
    queryset = Area.objects.all()
    serializer_class = AreaSerializer

    def get_queryset(self):
        qs = super().get_queryset()
        user = self.request.user
        if user.is_super_admin or user.is_staff:
            qs = qs.filter(organization=user.organization)
        else:
            my_team_ids = user.teams.values_list('id', flat=True)
            qs = qs.filter(teams__id__in=my_team_ids).distinct()
        return qs

    def perform_create(self, serializer):
        serializer.save(organization=self.request.user.organization)

    @action(detail=True, methods=['get'])
    def teams(self, request, pk=None):
        area = self.get_object()
        teams = area.teams.all()
        if not request.user.is_staff and not request.user.is_super_admin:
            teams = teams.filter(members=request.user)
        return Response(TeamSerializer(teams, many=True).data)


class TeamViewSet(viewsets.ModelViewSet):
    queryset = Team.objects.all()
    serializer_class = TeamSerializer

    def get_queryset(self):
        qs = super().get_queryset()
        user = self.request.user
        if user.is_super_admin or user.is_staff:
            qs = qs.filter(area__organization=user.organization)
        else:
            qs = qs.filter(members=user)
        area_id = self.request.query_params.get('area')
        if area_id:
            qs = qs.filter(area_id=area_id)
        return qs

    @action(detail=True, methods=['post'])
    def add_member(self, request, pk=None):
        team = self.get_object()
        user_id = request.data.get('user_id')
        try:
            user = User.objects.get(pk=user_id)
        except User.DoesNotExist:
            return Response({'error': 'Usuario no encontrado'}, status=status.HTTP_404_NOT_FOUND)
        team.members.add(user)
        return Response(TeamSerializer(team).data)

    @action(detail=True, methods=['post'])
    def remove_member(self, request, pk=None):
        team = self.get_object()
        user_id = request.data.get('user_id')
        try:
            user = User.objects.get(pk=user_id)
        except User.DoesNotExist:
            return Response({'error': 'Usuario no encontrado'}, status=status.HTTP_404_NOT_FOUND)
        team.members.remove(user)
        return Response(TeamSerializer(team).data)


class ProjectViewSet(viewsets.ModelViewSet):
    queryset = Project.objects.select_related('team', 'owner').prefetch_related(
        'sections__tasks__assignee',
        'sections__tasks__assigned_by',
        'sections__tasks__subtasks',
    ).all()

    def get_serializer_class(self):
        if self.action == 'list':
            return ProjectListSerializer
        return ProjectSerializer

    def get_queryset(self):
        qs = super().get_queryset()
        user = self.request.user
        if user.is_super_admin or user.is_staff:
            qs = qs.filter(team__area__organization=user.organization)
        else:
            qs = qs.filter(team__members=user)
        team_id = self.request.query_params.get('team')
        if team_id:
            qs = qs.filter(team_id=team_id)
        return qs


class SectionViewSet(viewsets.ModelViewSet):
    queryset = Section.objects.all()
    serializer_class = SectionSerializer

    def get_queryset(self):
        qs = super().get_queryset()
        project_id = self.request.query_params.get('project')
        if project_id:
            qs = qs.filter(project_id=project_id)
        return qs

    @action(detail=False, methods=['post'])
    def reorder(self, request):
        items = request.data.get('order', [])
        for item in items:
            Section.objects.filter(pk=item['id']).update(order=item['order'])
        return Response({'ok': True})


class TaskViewSet(viewsets.ModelViewSet):
    queryset = Task.objects.select_related('assignee', 'assigned_by', 'section__project').all()
    serializer_class = TaskSerializer

    def get_queryset(self):
        from django.db.models import Q
        qs = super().get_queryset()
        user = self.request.user

        if not user.is_staff:
            qs = qs.filter(
                Q(visibility='public') |
                Q(assignee=user) |
                Q(assigned_by=user)
            )

        section_id = self.request.query_params.get('section')
        project_id = self.request.query_params.get('project')
        assignee_id = self.request.query_params.get('assignee')
        status_filter = self.request.query_params.get('status')
        parent = self.request.query_params.get('parent')

        if section_id:
            qs = qs.filter(section_id=section_id)
        if project_id:
            qs = qs.filter(section__project_id=project_id)
        if assignee_id:
            qs = qs.filter(assignee_id=assignee_id)
        assigned_by_id = self.request.query_params.get('assigned_by')
        if assigned_by_id:
            qs = qs.filter(assigned_by_id=assigned_by_id)
        if status_filter:
            qs = qs.filter(status=status_filter)
        if parent == 'none':
            qs = qs.filter(parent__isnull=True)
        elif parent:
            qs = qs.filter(parent_id=parent)

        return qs.distinct()

    def perform_update(self, serializer):
        instance = serializer.instance
        user = self.request.user
        new_assignee = serializer.validated_data.get('assignee', instance.assignee)
        new_status = serializer.validated_data.get('status', instance.status)

        extra = {}
        assignee_changed = new_assignee and new_assignee != instance.assignee
        is_direct = False
        if assignee_changed:
            extra['assigned_by'] = user
            if user.is_staff:
                extra['assignment_status'] = Task.AssignmentStatus.DIRECT
                is_direct = True
            else:
                extra['assignment_status'] = Task.AssignmentStatus.PENDING

        if new_status == Task.Status.COMPLETED and instance.status != Task.Status.COMPLETED:
            extra['completed_at'] = timezone.now()
            task = serializer.save(**extra)
            if instance.recurrence_type != Task.RecurrenceType.NONE:
                instance.refresh_from_db()
                instance.create_next_occurrence()
        elif new_status != Task.Status.COMPLETED and instance.status == Task.Status.COMPLETED:
            extra['completed_at'] = None
            task = serializer.save(**extra)
        else:
            task = serializer.save(**extra)

        if assignee_changed:
            task.refresh_from_db()
            if is_direct:
                emails.notify_task_assigned_direct(task, user)
                notif.notify_task_assigned_direct(task, user)
            else:
                emails.notify_task_assigned(task, user)
                notif.notify_task_assigned(task, user)

    def perform_create(self, serializer):
        user = self.request.user
        assignee = serializer.validated_data.get('assignee')
        if assignee and assignee != user:
            if user.is_staff:
                task = serializer.save(assigned_by=user, assignment_status=Task.AssignmentStatus.DIRECT)
                emails.notify_task_assigned_direct(task, user)
                notif.notify_task_assigned_direct(task, user)
            else:
                task = serializer.save(assigned_by=user, assignment_status=Task.AssignmentStatus.PENDING)
                emails.notify_task_assigned(task, user)
                notif.notify_task_assigned(task, user)
        else:
            serializer.save(assigned_by=user, assignment_status=Task.AssignmentStatus.DIRECT)

    @action(detail=True, methods=['post'])
    def accept(self, request, pk=None):
        task = self.get_object()
        if task.assignee != request.user:
            return Response({'error': 'Solo el asignado puede aceptar'}, status=status.HTTP_403_FORBIDDEN)
        if task.assignment_status != Task.AssignmentStatus.PENDING:
            return Response({'error': 'Esta tarea no está pendiente de aprobación'}, status=status.HTTP_400_BAD_REQUEST)
        task.assignment_status = Task.AssignmentStatus.ACCEPTED
        task.save()
        emails.notify_assignment_accepted(task)
        notif.notify_assignment_accepted(task)
        return Response(TaskSerializer(task).data)

    @action(detail=True, methods=['post'])
    def reject(self, request, pk=None):
        task = self.get_object()
        if task.assignee != request.user:
            return Response({'error': 'Solo el asignado puede rechazar'}, status=status.HTTP_403_FORBIDDEN)
        if task.assignment_status != Task.AssignmentStatus.PENDING:
            return Response({'error': 'Esta tarea no está pendiente de aprobación'}, status=status.HTTP_400_BAD_REQUEST)
        # Guardamos el nombre antes de limpiar el assignee
        task._rejected_by_name = request.user.get_full_name() or request.user.username
        task.assignment_status = Task.AssignmentStatus.REJECTED
        task.assignee = None
        task.save()
        emails.notify_assignment_rejected(task)
        notif.notify_assignment_rejected(task)
        return Response(TaskSerializer(task).data)

    @action(detail=False, methods=['get'])
    def pending_assignments(self, request):
        tasks = Task.objects.filter(
            assignee=request.user,
            assignment_status=Task.AssignmentStatus.PENDING,
        )
        return Response(TaskSerializer(tasks, many=True).data)

    @action(detail=True, methods=['get'])
    def subtasks(self, request, pk=None):
        task = self.get_object()
        subtasks = task.subtasks.all()
        return Response(TaskSerializer(subtasks, many=True).data)

    @action(detail=False, methods=['post'])
    def reorder(self, request):
        items = request.data.get('order', [])
        for item in items:
            Task.objects.filter(pk=item['id']).update(
                order=item['order'],
                section_id=item.get('section', None) or Task.objects.get(pk=item['id']).section_id,
            )
        return Response({'ok': True})


class AssetViewSet(viewsets.ModelViewSet):
    queryset = Asset.objects.all()
    serializer_class = AssetSerializer

    def get_queryset(self):
        qs = super().get_queryset()
        project_id = self.request.query_params.get('project')
        task_id = self.request.query_params.get('task')
        category = self.request.query_params.get('category')
        if project_id:
            qs = qs.filter(project_id=project_id)
        if task_id:
            qs = qs.filter(task_id=task_id)
        if category:
            if category == 'email':
                qs = qs.filter(asset_type=Asset.AssetType.EMAIL)
            elif category == 'link':
                qs = qs.filter(asset_type=Asset.AssetType.LINK)
            else:
                qs = qs.filter(asset_type=Asset.AssetType.FILE).exclude(mime_type='message/rfc822')
        return qs

    def perform_create(self, serializer):
        asset = serializer.save(uploaded_by=self.request.user)
        if asset.file:
            asset.file_size = asset.file.size
            import mimetypes
            mime, _ = mimetypes.guess_type(asset.file.name)
            asset.mime_type = mime or ''
            if asset.file_extension == 'eml':
                asset.asset_type = Asset.AssetType.EMAIL
                asset.mime_type = 'message/rfc822'
            if not asset.name or asset.name == 'undefined':
                asset.name = asset.file.name.split('/')[-1]
            asset.save()

    @action(detail=True, methods=['get'])
    def parse_eml(self, request, pk=None):
        import email as email_lib
        from email import policy
        import base64

        asset = self.get_object()
        if asset.file_extension != 'eml' and asset.asset_type != Asset.AssetType.EMAIL:
            return Response({'error': 'No es un archivo .eml'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            with asset.file.open('rb') as f:
                msg = email_lib.message_from_binary_file(f, policy=policy.default)

            result = {
                'subject': str(msg.get('subject', '')),
                'from': str(msg.get('from', '')),
                'to': str(msg.get('to', '')),
                'cc': str(msg.get('cc', '')) or None,
                'date': str(msg.get('date', '')),
                'body_text': '',
                'body_html': '',
                'attachments': [],
            }

            def decode_payload(part):
                payload = part.get_payload(decode=True)
                if not payload:
                    return ''
                charset = part.get_content_charset() or 'utf-8'
                for enc in [charset, 'utf-8', 'latin-1', 'windows-1252']:
                    try:
                        return payload.decode(enc)
                    except (UnicodeDecodeError, LookupError):
                        continue
                return payload.decode('utf-8', errors='replace')

            att_index = 0
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    disposition = str(part.get('Content-Disposition', ''))

                    if 'attachment' in disposition:
                        filename = part.get_filename() or 'adjunto'
                        payload = part.get_payload(decode=True)
                        result['attachments'].append({
                            'index': att_index,
                            'filename': filename,
                            'content_type': content_type,
                            'size': len(payload) if payload else 0,
                        })
                        att_index += 1
                    elif content_type == 'text/plain' and not result['body_text']:
                        result['body_text'] = decode_payload(part)
                    elif content_type == 'text/html' and not result['body_html']:
                        result['body_html'] = decode_payload(part)
            else:
                text = decode_payload(msg)
                if msg.get_content_type() == 'text/html':
                    result['body_html'] = text
                else:
                    result['body_text'] = text

            return Response(result)
        except Exception as e:
            return Response({'error': f'Error al parsear el archivo: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['get'], url_path='eml_attachment/(?P<att_index>[0-9]+)')
    def eml_attachment(self, request, pk=None, att_index=None):
        import email as email_lib
        from email import policy
        from django.http import HttpResponse

        asset = self.get_object()
        if asset.file_extension != 'eml' and asset.asset_type != Asset.AssetType.EMAIL:
            return Response({'error': 'No es un archivo .eml'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            att_index = int(att_index)
            with asset.file.open('rb') as f:
                msg = email_lib.message_from_binary_file(f, policy=policy.default)

            current = 0
            for part in msg.walk():
                disposition = str(part.get('Content-Disposition', ''))
                if 'attachment' in disposition:
                    if current == att_index:
                        payload = part.get_payload(decode=True)
                        filename = part.get_filename() or 'adjunto'
                        content_type = part.get_content_type()
                        response = HttpResponse(payload, content_type=content_type)
                        response['Content-Disposition'] = f'attachment; filename="{filename}"'
                        return response
                    current += 1

            return Response({'error': 'Adjunto no encontrado'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['get'])
    def preview(self, request, pk=None):
        from django.http import HttpResponse
        asset = self.get_object()
        if not asset.file:
            return Response({'error': 'No hay archivo'}, status=status.HTTP_400_BAD_REQUEST)

        ext = asset.file_extension
        try:
            if ext in ('jpg', 'jpeg', 'png', 'gif', 'webp', 'svg', 'bmp'):
                return Response({'type': 'image', 'url': asset.file.url})

            if ext == 'pdf':
                return Response({'type': 'pdf', 'url': asset.file.url})

            if ext in ('txt', 'csv', 'log', 'md', 'json', 'xml', 'html', 'css', 'js', 'py'):
                with asset.file.open('rb') as f:
                    raw = f.read()
                for enc in ['utf-8', 'latin-1', 'windows-1252']:
                    try:
                        content = raw.decode(enc)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    content = raw.decode('utf-8', errors='replace')
                return Response({'type': 'text', 'content': content, 'filename': asset.name})

            if ext in ('xlsx', 'xls'):
                import openpyxl
                wb = openpyxl.load_workbook(asset.file.path, read_only=True, data_only=True)
                sheets = []
                for ws in wb.worksheets:
                    rows = []
                    for row in ws.iter_rows(max_row=200, values_only=True):
                        rows.append([str(cell) if cell is not None else '' for cell in row])
                    sheets.append({'name': ws.title, 'rows': rows})
                wb.close()
                return Response({'type': 'excel', 'sheets': sheets, 'filename': asset.name})

            if ext in ('docx',):
                import docx
                from html import escape
                doc = docx.Document(asset.file.path)
                html_parts = ['<meta charset="utf-8">']
                for para in doc.paragraphs:
                    style = para.style.name if para.style else ''
                    text = escape(para.text)
                    if not text.strip():
                        html_parts.append('<br>')
                    elif 'Heading 1' in style:
                        html_parts.append(f'<h1>{text}</h1>')
                    elif 'Heading 2' in style:
                        html_parts.append(f'<h2>{text}</h2>')
                    elif 'Heading 3' in style:
                        html_parts.append(f'<h3>{text}</h3>')
                    else:
                        html_parts.append(f'<p>{text}</p>')
                for table in doc.tables:
                    html_parts.append('<table border="1" cellpadding="6" cellspacing="0" style="border-collapse:collapse;width:100%">')
                    for row in table.rows:
                        html_parts.append('<tr>')
                        for cell in row.cells:
                            html_parts.append(f'<td>{escape(cell.text)}</td>')
                        html_parts.append('</tr>')
                    html_parts.append('</table>')
                return Response({'type': 'word', 'html': '\n'.join(html_parts), 'filename': asset.name}, content_type='application/json; charset=utf-8')

            return Response({'type': 'unsupported', 'message': 'Vista previa no disponible para este tipo de archivo'})
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)


class PlanViewSet(viewsets.ModelViewSet):
    queryset = Plan.objects.all()
    serializer_class = PlanSerializer
    permission_classes = [IsSuperAdmin]


class OrganizationViewSet(viewsets.ModelViewSet):
    queryset = Organization.objects.all()
    serializer_class = OrganizationSerializer
    permission_classes = [IsSuperAdmin]

    @action(detail=True, methods=['get'])
    def users(self, request, pk=None):
        org = self.get_object()
        users = org.members.all()
        return Response(UserSerializer(users, many=True).data)

    @action(detail=True, methods=['post'])
    def add_user(self, request, pk=None):
        org = self.get_object()
        if org.members.filter(is_active=True).count() >= org.max_users:
            return Response({'error': f'Límite de {org.max_users} usuarios alcanzado'}, status=status.HTTP_400_BAD_REQUEST)
        user_data = request.data
        password = user_data.pop('password', 'changeme123')
        user = User.objects.create(**user_data, organization=org)
        user.set_password(password)
        user.save()
        return Response(UserSerializer(user).data, status=status.HTTP_201_CREATED)

    @action(detail=True, methods=['post'])
    def toggle_active(self, request, pk=None):
        org = self.get_object()
        org.is_active = not org.is_active
        org.save()
        return Response(OrganizationSerializer(org).data)


class NotificationViewSet(viewsets.ReadOnlyModelViewSet):
    serializer_class = NotificationSerializer

    def get_queryset(self):
        qs = Notification.objects.filter(recipient=self.request.user).select_related('actor', 'task__section')
        if self.request.query_params.get('unread') == 'true':
            qs = qs.filter(is_read=False)
        return qs

    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset()[:50]
        serializer = self.get_serializer(queryset, many=True)
        return Response(serializer.data)

    @action(detail=False, methods=['get'])
    def unread_count(self, request):
        count = Notification.objects.filter(recipient=request.user, is_read=False).count()
        return Response({'count': count})

    @action(detail=True, methods=['post'])
    def mark_read(self, request, pk=None):
        notification = self.get_queryset().filter(pk=pk).first()
        if not notification:
            return Response({'error': 'No encontrada'}, status=status.HTTP_404_NOT_FOUND)
        notification.is_read = True
        notification.save()
        return Response(NotificationSerializer(notification).data)

    @action(detail=False, methods=['post'])
    def mark_all_read(self, request):
        Notification.objects.filter(recipient=request.user, is_read=False).update(is_read=True)
        return Response({'ok': True})
